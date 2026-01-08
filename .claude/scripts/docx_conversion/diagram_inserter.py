"""
附图插入器 - 将 Mermaid 图表插入到 DOCX 文档

本脚本负责：
1. 从 Markdown 中提取 Mermaid 代码块
2. 使用 mermaid-cli 渲染为 PNG 图片
3. 将图片插入到 DOCX 对应位置
"""

import re
import json
import subprocess
import argparse
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print("错误: 未安装 python-docx 库")
    print("请运行: pip install python-docx")
    exit(1)


class DiagramInserter:
    """附图插入器"""

    # 图片-章节映射表
    DIAGRAM_SECTION_MAP = {
        "图1": {"section": "4.（2）技术方案", "position": "first_paragraph"},
        "图2": {"section": "4.（2）技术方案", "position": "after_keyword", "keyword": "软件架构"},
        "图3": {"section": "4.（2）技术方案", "position": "after_keyword", "keyword": "功能模块"},
        "图4": {"section": "5. 具体实施方式", "position": "before_keyword", "keyword": "设备发现"},
        "图5": {"section": "5. 具体实施方式", "position": "before_keyword", "keyword": "IP分配"},
        "图6": {"section": "5. 具体实施方式", "position": "before_keyword", "keyword": "时序|交互|消息"},
        "图7": {"section": "4.（2）技术方案", "position": "after_keyword", "keyword": "冲突检测"},
        "图8": {"section": "5. 具体实施方式", "position": "before_keyword", "keyword": "初始化"},
        "图9": {"section": "5. 具体实施方式", "position": "before_keyword", "keyword": "故障切换"},
        "图10": {"section": "5. 具体实施方式", "position": "before_keyword", "keyword": "协议格式|设备发现协议"},
        "图11": {"section": "5. 具体实施方式", "position": "before_keyword", "keyword": "协议格式|IP分配协议"},
        "图12": {"section": "4.（3）有益效果", "position": "first_paragraph"},
    }

    def __init__(self, markdown_path: str, docx_path: str,
                 diagram_desc_path: str, images_output_dir: str):
        """
        初始化插入器

        Args:
            markdown_path: Markdown 文件路径
            docx_path: DOCX 文件路径
            diagram_desc_path: 附图说明文件路径
            images_output_dir: 图片输出目录
        """
        self.markdown_path = Path(markdown_path)
        self.docx_path = Path(docx_path)
        self.diagram_desc_path = Path(diagram_desc_path)
        self.images_output_dir = Path(images_output_dir)
        self.images_output_dir.mkdir(parents=True, exist_ok=True)

    def check_mermaid_cli(self) -> bool:
        """检查 mermaid-cli 是否已安装"""
        try:
            result = subprocess.run(
                ['mmdc', '--version'],
                check=True,
                capture_output=True,
                text=True
            )
            return True
        except (subprocess.CalledProcessError, FileNotFoundError):
            return False

    def extract_mermaid_from_markdown(self) -> List[Dict[str, Any]]:
        """
        从 Markdown 中提取 Mermaid 代码块

        Returns:
            附图信息列表
        """
        content = self.markdown_path.read_text(encoding='utf-8')

        # 匹配附图块格式：**附图X：名称** \n ```mermaid\n ... ``` \n 图X说明：...
        pattern = re.compile(
            r'\*\*附图(图\d+)：(.+?)\*\*\s*\n\s*```mermaid\n(.*?)```\s*\n\s*图\d+说明：(.*?)(?=\n\n|\n\*\*|$)',
            re.DOTALL
        )

        diagrams = []
        for match in pattern.finditer(content):
            diagrams.append({
                'number': match.group(1).strip(),
                'name': match.group(2).strip(),
                'mermaid_code': match.group(3).strip(),
                'description': match.group(4).strip()
            })

        # 如果没有找到，尝试从附图说明文件中提取
        if not diagrams and self.diagram_desc_path.exists():
            diagrams = self._extract_from_diagram_desc()

        return diagrams

    def _extract_from_diagram_desc(self) -> List[Dict[str, Any]]:
        """从附图说明文件中提取 Mermaid 代码"""
        content = self.diagram_desc_path.read_text(encoding='utf-8')

        diagrams = []
        current_diagram = None

        for line in content.split('\n'):
            # 匹配图标题
            if line.startswith('### 图'):
                if current_diagram:
                    diagrams.append(current_diagram)
                current_diagram = {'number': '', 'name': '', 'mermaid_code': '', 'description': ''}
                # 提取图编号和名称
                match = re.match(r'### (图\d+): (.+)', line)
                if match:
                    current_diagram['number'] = match.group(1)
                    current_diagram['name'] = match.group(2)

            # 提取 Mermaid 代码
            elif line.startswith('```mermaid') and current_diagram:
                current_diagram['mermaid_code'] = ''
            elif line.startswith('```') and current_diagram and current_diagram['mermaid_code'] != '':
                pass  # 结束代码块
            elif current_diagram and current_diagram.get('mermaid_code') is not None:
                if current_diagram['mermaid_code'] == '':
                    # 代码开始后的第一行
                    pass
                elif 'mermaid_code' in current_diagram:
                    current_diagram['mermaid_code'] += line + '\n'

            # 提取附图说明
            elif '附图说明' in line and current_diagram:
                match = re.search(r'本图展示了(.+?)，与说明书中的(.+?)部分相对应', line)
                if match:
                    current_diagram['description'] = match.group(1)

        if current_diagram:
            diagrams.append(current_diagram)

        return diagrams

    def render_mermaid_to_image(self, mermaid_code: str,
                                output_path: Path) -> bool:
        """
        使用 mermaid-cli 渲染 Mermaid 代码为图片

        Args:
            mermaid_code: Mermaid 代码
            output_path: 输出图片路径

        Returns:
            是否渲染成功
        """
        # 创建临时 mermaid 文件
        temp_mmd = self.images_output_dir / "temp.mmd"
        temp_mmd.write_text(mermaid_code, encoding='utf-8')

        try:
            # 调用 mermaid CLI
            subprocess.run([
                'mmdc', '-i', str(temp_mmd), '-o', str(output_path),
                '-b', 'transparent', '-s', '2'  # 2倍缩放以提高清晰度
            ], check=True, capture_output=True)

            # 清理临时文件
            temp_mmd.unlink()
            return True
        except subprocess.CalledProcessError as e:
            print(f"渲染失败: {e.stderr.decode('utf-8', errors='ignore') if e.stderr else e}")
            temp_mmd.unlink()
            return False

    def find_insertion_position(self, doc: Document,
                                diagram_number: str) -> Tuple[int, str]:
        """
        在 DOCX 中查找最佳插入位置

        Args:
            doc: DOCX 文档对象
            diagram_number: 附图编号（如"图1"）

        Returns:
            (段落索引, 目标章节)
        """
        # 获取映射信息
        mapping = self.DIAGRAM_SECTION_MAP.get(diagram_number, {})
        target_section = mapping.get("section", "4.（2）技术方案")
        position_type = mapping.get("position", "first_paragraph")

        # 查找目标章节
        section_idx = self._find_section_index(doc, target_section)
        if section_idx == -1:
            print(f"警告: 未找到章节 {target_section}，使用默认位置")
            return (0, target_section)

        # 根据位置类型确定插入位置
        if position_type == "first_paragraph":
            # 章节第一个段落前
            return (section_idx, target_section)

        elif position_type == "before_keyword":
            # 关键词段落前
            keyword = mapping.get("keyword", "")
            if keyword:
                idx = self._find_keyword_before(doc, section_idx, keyword)
                if idx != -1:
                    return (idx, target_section)
            return (section_idx, target_section)

        elif position_type == "after_keyword":
            # 关键词段落后
            keyword = mapping.get("keyword", "")
            if keyword:
                idx = self._find_keyword_after(doc, section_idx, keyword)
                if idx != -1:
                    return (idx, target_section)
            return (section_idx, target_section)

        return (section_idx, target_section)

    def _find_section_index(self, doc: Document, section_name: str) -> int:
        """查找章节索引"""
        for i, para in enumerate(doc.paragraphs):
            if section_name in para.text:
                return i
        return -1

    def _find_keyword_before(self, doc: Document, start_idx: int, keyword: str) -> int:
        """在章节中查找包含关键词的第一个段落"""
        pattern = re.compile(keyword)
        for i in range(start_idx, min(start_idx + 50, len(doc.paragraphs))):
            if pattern.search(doc.paragraphs[i].text):
                return i
        return -1

    def _find_keyword_after(self, doc: Document, start_idx: int, keyword: str) -> int:
        """在章节中查找包含关键词的最后一个段落"""
        pattern = re.compile(keyword)
        last_match = -1
        for i in range(start_idx, min(start_idx + 50, len(doc.paragraphs))):
            if pattern.search(doc.paragraphs[i].text):
                last_match = i
        return last_match if last_match != -1 else start_idx

    def insert_diagram_into_docx(self, doc: Document,
                                  diagram_info: Dict[str, Any]) -> bool:
        """
        将附图插入到 DOCX 文档

        Args:
            doc: DOCX 文档对象
            diagram_info: 附图信息字典

        Returns:
            是否成功
        """
        # 渲染图片
        safe_name = re.sub(r'[<>:"/\\|?*]', '_', diagram_info['name'])
        image_filename = f"{diagram_info['number']}_{safe_name}.png"
        image_path = self.images_output_dir / image_filename

        if not self.render_mermaid_to_image(diagram_info['mermaid_code'], image_path):
            print(f"跳过 {diagram_info['number']}：渲染失败")
            return False

        # 查找插入位置
        insert_idx, target_section = self.find_insertion_position(
            doc, diagram_info['number']
        )

        # 插入图片
        try:
            # 获取插入位置的段落元素
            target_para = doc.paragraphs[insert_idx]._element

            # 创建新的图片段落
            new_para = OxmlElement('w:p')
            doc.paragraphs[insert_idx]._element.addnext(new_para)
            new_para_obj = Document(new_para.getparent()).paragraphs[-1]

            # 添加标题
            title_run = new_para_obj.add_run(f"\n附图{diagram_info['number']}：{diagram_info['name']}\n")
            title_run.font.size = Pt(11)
            title_run.font.bold = True
            title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '思源黑体 CN Bold')

            # 添加图片
            try:
                new_para_obj.add_run().add_picture(str(image_path), width=Inches(5.0))
            except Exception as e:
                print(f"插入图片失败: {e}")
                return False

            # 添加说明段落
            desc_para = OxmlElement('w:p')
            new_para._element.addnext(desc_para)
            desc_para_obj = Document(desc_para.getparent()).paragraphs[-1]
            desc_run = desc_para_obj.add_run(f"图{diagram_info['number']}说明：{diagram_info['description']}")
            desc_run.font.size = Pt(9)
            desc_run.font.italic = True
            desc_run._element.rPr.rFonts.set(qn('w:eastAsia'), '思源黑体 CN Normal')

            diagram_info['target_section'] = target_section
            diagram_info['image_path'] = str(image_path)
            return True

        except Exception as e:
            print(f"插入 {diagram_info['number']} 失败: {e}")
            return False

    def insert_all_diagrams(self) -> Dict[str, Any]:
        """
        插入所有附图

        Returns:
            插入报告
        """
        # 检查 mermaid-cli
        if not self.check_mermaid_cli():
            return {
                'success': False,
                'error': 'mermaid_cli_not_installed',
                'message': 'mermaid-cli 未安装，请运行: npm install -g @mermaid-js/mermaid-cli'
            }

        # 提取 Mermaid 代码
        diagrams = self.extract_mermaid_from_markdown()

        if not diagrams:
            return {
                'success': False,
                'error': 'no_diagrams_found',
                'message': '未找到附图，请确认 diagram-generator 已成功执行'
            }

        # 加载 DOCX
        doc = Document(str(self.docx_path))

        # 插入每幅图
        insertions = []
        success_count = 0

        for diagram_info in diagrams:
            diagram_info['status'] = 'pending'
            if self.insert_diagram_into_docx(doc, diagram_info):
                diagram_info['status'] = 'success'
                success_count += 1
            else:
                diagram_info['status'] = 'failed'
                diagram_info['error'] = '插入或渲染失败'

            insertions.append({
                'diagram_number': diagram_info['number'],
                'diagram_name': diagram_info['name'],
                'target_section': diagram_info.get('target_section', '未知'),
                'image_path': diagram_info.get('image_path', ''),
                'status': diagram_info['status']
            })

        # 保存 DOCX
        try:
            doc.save(str(self.docx_path))
            docx_modified = True
        except Exception as e:
            docx_modified = False
            print(f"保存 DOCX 失败: {e}")

        # 生成报告
        report = {
            'success': docx_modified,
            'insertion_timestamp': datetime.now().isoformat(),
            'total_diagrams': len(diagrams),
            'success_count': success_count,
            'insertions': insertions,
            'docx_modified': docx_modified
        }

        return report


def main():
    """主函数"""
    parser = argparse.ArgumentParser(description='附图插入器 - 将 Mermaid 图表插入到 DOCX')
    parser.add_argument('markdown_path', help='Markdown 文件路径')
    parser.add_argument('docx_path', help='DOCX 文件路径')
    parser.add_argument('diagram_desc_path', help='附图说明文件路径')
    parser.add_argument('images_output_dir', help='图片输出目录')
    parser.add_argument('--report', help='报告输出路径（JSON格式）', default=None)

    args = parser.parse_args()

    # 创建插入器
    inserter = DiagramInserter(
        args.markdown_path,
        args.docx_path,
        args.diagram_desc_path,
        args.images_output_dir
    )

    # 执行插入
    report = inserter.insert_all_diagrams()

    # 输出报告
    print(json.dumps(report, ensure_ascii=False, indent=2))

    # 保存报告到文件
    if args.report:
        report_path = Path(args.report)
        report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding='utf-8')
        print(f"\n报告已保存到: {report_path}")

    return 0 if report['success'] else 1


if __name__ == '__main__':
    exit(main())
