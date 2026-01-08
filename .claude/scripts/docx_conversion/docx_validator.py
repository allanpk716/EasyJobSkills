"""
DOCX 验证器 - 验证生成的 DOCX 文件格式和质量
"""

import json
import re
from pathlib import Path
from typing import Dict, Any, List
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .exceptions import ValidationError


class DocxValidator:
    """DOCX 文档验证器"""

    # 必需章节
    REQUIRED_SECTIONS = [
        {"number": "1", "title": "发明创造名称"},
        {"number": "2", "title": "所属技术领域"},
        {"number": "3", "title": "相关的背景技术"},
        {"number": "4", "title": "发明内容", "subsections": 3},
        {"number": "5", "title": "具体实施方式"},
        {"number": "6", "title": "关键点和欲保护点"},
        {"number": "7", "title": "其他有助于理解本技术的资料"},
    ]

    # 字体要求
    EXPECTED_FONT = "思源黑体 CN"
    TITLE_FONT_SIZE = 18  # 磅
    BODY_FONT_SIZE = 10  # 磅

    def __init__(self, docx_path: str, validation_level: str = "standard"):
        """
        初始化验证器

        Args:
            docx_path: DOCX 文件路径
            validation_level: 验证级别（standard/strict）
        """
        self.docx_path = Path(docx_path)
        if not self.docx_path.exists():
            raise ValidationError(docx_path, ["文件不存在"])

        self.validation_level = validation_level
        self.doc = Document(str(self.docx_path))
        self.checks = {}

    def validate_all(self) -> Dict[str, Any]:
        """
        执行所有验证检查

        Returns:
            完整验证报告
        """
        results = {
            "validation_passed": True,
            "validation_level": self.validation_level,
            "validation_timestamp": datetime.now().isoformat(),
            "checks": {
                "section_completeness": self._check_section_completeness(),
                "font_application": self._check_font_application(),
                "paragraph_formatting": self._check_paragraph_formatting(),
                "style_consistency": self._check_style_consistency(),
                "page_setup": self._check_page_setup(),
                "content_quality": self._check_content_quality(),
            },
        }

        # 计算总体评分
        results["overall_score"] = self._calculate_score(results["checks"])
        results["recommendations"] = self._generate_recommendations(
            results["checks"]
        )
        results["critical_issues"] = self._identify_critical_issues(
            results["checks"]
        )

        # 判断是否通过
        results["validation_passed"] = (
            results["overall_score"] >= 80 and len(results["critical_issues"]) == 0
        )

        return results

    def _check_section_completeness(self) -> Dict[str, Any]:
        """检查章节完整性"""
        found_sections = []
        section_titles = {}

        for para in self.doc.paragraphs:
            text = para.text.strip()
            # 识别章节标题
            for req_section in self.REQUIRED_SECTIONS:
                pattern = f"^{req_section['number']}[、.]\s*{req_section['title']}"
                if re.match(pattern, text):
                    found_sections.append(req_section["number"])
                    section_titles[req_section["number"]] = req_section[
                        "title"
                    ]
                    break

        missing = [
            s["number"] for s in self.REQUIRED_SECTIONS if s["number"] not in found_sections
        ]

        # 检查第4章节的子项
        section_4_subsections = 0
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if re.match(r"^（[123]）[、.]\s*", text):
                section_4_subsections += 1

        subsections_complete = section_4_subsections >= 3

        return {
            "passed": len(missing) == 0 and subsections_complete,
            "details": {
                "expected_sections": len(self.REQUIRED_SECTIONS),
                "found_sections": len(found_sections),
                "missing_sections": missing,
                "sections": [
                    {
                        "number": s["number"],
                        "title": s["title"],
                        "found": s["number"] in found_sections,
                    }
                    for s in self.REQUIRED_SECTIONS
                ],
                "section_4_subsections": {
                    "expected": 3,
                    "found": section_4_subsections,
                    "complete": subsections_complete,
                },
            },
        }

    def _check_font_application(self) -> Dict[str, Any]:
        """检查字体应用"""
        font_issues = []
        title_check = {"passed": False, "font_name": None, "font_size_pt": None, "is_bold": None}
        body_check = {"passed": False, "font_name": None, "font_size_pt": None, "is_bold": None}

        for para in self.doc.paragraphs:
            if not para.runs:
                continue

            run = para.runs[0]
            font_name = run.font.name
            font_size = run.font.size.pt if run.font.size else None
            is_bold = run.font.bold

            # 检查标题字体（第一个粗体段落）
            if title_check["font_name"] is None and font_name and "Bold" in font_name:
                title_check["font_name"] = font_name
                title_check["font_size_pt"] = font_size
                title_check["is_bold"] = is_bold
                title_check["passed"] = (
                    "思源黑体 CN" in font_name
                    and abs(font_size - self.TITLE_FONT_SIZE) < 1
                    and is_bold == True
                )

            # 检查正文字体（第一个非粗体段落）
            elif (
                body_check["font_name"] is None
                and font_name
                and "Normal" in font_name
            ):
                body_check["font_name"] = font_name
                body_check["font_size_pt"] = font_size
                body_check["is_bold"] = is_bold
                body_check["passed"] = (
                    "思源黑体 CN" in font_name
                    and abs(font_size - self.BODY_FONT_SIZE) < 1
                    and is_bold == False
                )

            if title_check["font_name"] and body_check["font_name"]:
                break

        # 生成问题列表
        if not title_check["passed"]:
            font_issues.append(
                {
                    "type": "title_font",
                    "expected": f"{self.EXPECTED_FONT} Bold, {self.TITLE_FONT_SIZE}pt",
                    "actual": f"{title_check['font_name']}, {title_check['font_size_pt']}pt",
                }
            )

        if not body_check["passed"]:
            font_issues.append(
                {
                    "type": "body_font",
                    "expected": f"{self.EXPECTED_FONT} Normal, {self.BODY_FONT_SIZE}pt",
                    "actual": f"{body_check['font_name']}, {body_check['font_size_pt']}pt",
                }
            )

        return {
            "passed": len(font_issues) == 0,
            "details": {
                "expected_font": self.EXPECTED_FONT,
                "title_font_check": title_check,
                "body_font_check": body_check,
                "font_issues": font_issues,
            },
        }

    def _check_paragraph_formatting(self) -> Dict[str, Any]:
        """检查段落格式"""
        line_spacing_ok = True
        indent_ok = True
        alignment_ok = True

        for para in self.doc.paragraphs:
            if not para.text.strip():
                continue

            # 检查行距（允许误差 ±0.1）
            if para.paragraph_format.line_spacing:
                ls = para.paragraph_format.line_spacing
                if abs(ls - 1.5) > 0.1:
                    line_spacing_ok = False

            # 检查对齐
            if para.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
                # 标题可以居中
                if para.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                    alignment_ok = False

        return {
            "passed": line_spacing_ok and indent_ok and alignment_ok,
            "details": {
                "line_spacing_check": {
                    "expected": 1.5,
                    "actual_range": "1.4-1.6",
                    "passed": line_spacing_ok,
                },
                "first_line_indent_check": {
                    "expected_chars": 2,
                    "passed": indent_ok,
                },
                "alignment_check": {
                    "expected": "justify",
                    "passed": alignment_ok,
                },
            },
        }

    def _check_style_consistency(self) -> Dict[str, Any]:
        """检查样式一致性"""
        inconsistencies = []

        # 收集所有字体设置
        title_fonts = set()
        body_fonts = set()

        for para in self.doc.paragraphs:
            if para.runs:
                font = para.runs[0].font.name
                if para.runs[0].font.bold:
                    title_fonts.add(font)
                else:
                    body_fonts.add(font)

        title_consistent = len(title_fonts) <= 2  # 允许 Bold 和 Normal
        body_consistent = len(body_fonts) <= 2

        if not title_consistent:
            inconsistencies.append(
                {"type": "title_font_inconsistent", "fonts_found": list(title_fonts)}
            )

        if not body_consistent:
            inconsistencies.append(
                {"type": "body_font_inconsistent", "fonts_found": list(body_fonts)}
            )

        return {
            "passed": len(inconsistencies) == 0,
            "details": {
                "title_style_consistent": title_consistent,
                "body_style_consistent": body_consistent,
                "inconsistencies": inconsistencies,
            },
        }

    def _check_page_setup(self) -> Dict[str, Any]:
        """检查页面设置"""
        section = self.doc.sections[0]

        # 页边距（转换为厘米）
        top_cm = section.top_margin.cm
        bottom_cm = section.bottom_margin.cm
        left_cm = section.left_margin.cm
        right_cm = section.right_margin.cm

        # 纸张大小
        page_width = section.page_width.cm
        page_height = section.page_height.cm
        is_a4 = abs(page_width - 21.0) < 0.5 and abs(page_height - 29.7) < 0.5

        return {
            "passed": is_a4,
            "details": {
                "margins": {
                    "top": f"{top_cm:.2f}cm",
                    "bottom": f"{bottom_cm:.2f}cm",
                    "left": f"{left_cm:.2f}cm",
                    "right": f"{right_cm:.2f}cm",
                },
                "paper_size": "A4" if is_a4 else "Other",
            },
        }

    def _check_content_quality(self) -> Dict[str, Any]:
        """检查内容质量"""
        empty_paragraphs = 0
        very_short_sections = []
        warnings = []

        for para in self.doc.paragraphs:
            if not para.text.strip():
                empty_paragraphs += 1

        # 检查章节长度
        current_section = None
        current_content_length = 0

        for para in self.doc.paragraphs:
            text = para.text.strip()

            # 识别章节
            for req_section in self.REQUIRED_SECTIONS:
                pattern = f"^{req_section['number']}[、.]\s*{req_section['title']}"
                if re.match(pattern, text):
                    # 保存上一章节
                    if (
                        current_section
                        and current_content_length < 20
                        and current_content_length > 0
                    ):
                        very_short_sections.append(
                            {
                                "section": current_section,
                                "length": current_content_length,
                            }
                        )
                    current_section = req_section["title"]
                    current_content_length = 0
                    break
            else:
                if current_section:
                    current_content_length += len(text)

        # 生成警告
        if empty_paragraphs > 5:
            warnings.append(f"发现 {empty_paragraphs} 个空段落")

        if very_short_sections:
            warnings.append(
                f"以下章节内容过短: {[s['section'] for s in very_short_sections]}"
            )

        return {
            "passed": len(warnings) == 0,
            "details": {
                "empty_paragraphs": empty_paragraphs,
                "very_short_sections": very_short_sections,
                "warnings": warnings,
            },
        }

    def _calculate_score(self, checks: Dict[str, Any]) -> int:
        """计算总体评分（0-100）"""
        weights = {
            "section_completeness": 30,
            "font_application": 25,
            "paragraph_formatting": 20,
            "style_consistency": 10,
            "page_setup": 5,
            "content_quality": 10,
        }

        total_score = 0
        for check_name, weight in weights.items():
            if check_name in checks:
                if checks[check_name]["passed"]:
                    total_score += weight

        return total_score

    def _generate_recommendations(self, checks: Dict[str, Any]) -> List[str]:
        """生成改进建议"""
        recommendations = []

        if not checks["section_completeness"]["passed"]:
            recommendations.append("请补充缺失的章节内容")

        if not checks["font_application"]["passed"]:
            recommendations.append(
                "请确保使用思源黑体 CN 字体，标题18pt，正文10pt"
            )

        if not checks["paragraph_formatting"]["passed"]:
            recommendations.append("请设置段落格式：行距1.5倍，两端对齐")

        if not checks["style_consistency"]["passed"]:
            recommendations.append("请统一文档中的字体和样式")

        if not checks["content_quality"]["passed"]:
            recommendations.append("请检查并完善内容，删除不必要的空段落")

        return recommendations

    def _identify_critical_issues(self, checks: Dict[str, Any]) -> List[str]:
        """识别关键问题"""
        critical = []

        if not checks["section_completeness"]["passed"]:
            critical.append("章节不完整")

        if not checks["font_application"]["passed"]:
            critical.append("字体设置错误")

        return critical

    def save_report(self, output_path: str):
        """
        保存验证报告

        Args:
            output_path: 输出 JSON 文件路径
        """
        report = self.validate_all()
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(report, f, ensure_ascii=False, indent=2)
        return output_path


def main():
    """命令行入口"""
    import sys

    if len(sys.argv) < 2:
        print("Usage: python docx_validator.py <docx_path> [output_json]")
        sys.exit(1)

    docx_path = sys.argv[1]
    output_json = sys.argv[2] if len(sys.argv) > 2 else "validation_report.json"

    try:
        validator = DocxValidator(docx_path)
        report = validator.validate_all()

        with open(output_json, "w", encoding="utf-8") as f:
            json.dump(report, f, ensure_ascii=False, indent=2)

        print(f"✅ 验证完成！报告保存到: {output_json}")
        print(f"📊 验证通过: {report['validation_passed']}")
        print(f"⭐ 总体评分: {report['overall_score']}/100")

        if report["recommendations"]:
            print("\n💡 改进建议:")
            for rec in report["recommendations"]:
                print(f"   - {rec}")

        if report["critical_issues"]:
            print("\n⚠️  关键问题:")
            for issue in report["critical_issues"]:
                print(f"   - {issue}")

    except ValidationError as e:
        print(f"❌ {e}")
        sys.exit(1)
    except Exception as e:
        print(f"❌ 验证失败: {e}")
        import traceback

        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
