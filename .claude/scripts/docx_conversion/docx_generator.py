"""
DOCX 生成器 - 基于模板填充内容并设置格式
"""

import json
import re
from pathlib import Path
from typing import Dict, Any
from datetime import datetime

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .font_utils import FontChecker, FontInstaller
from .format_utils import ParagraphFormatter
from .exceptions import FontNotFoundError, TemplateNotFoundError


class DocxGenerator:
    """DOCX 文档生成器"""

    # 字体配置
    FONT_NAME = "思源黑体 CN"
    TITLE_FONT_NAME = "思源黑体 CN Bold"
    BODY_FONT_NAME = "思源黑体 CN Normal"

    TITLE_FONT_SIZE = 18  # 磅
    BODY_FONT_SIZE = 10  # 磅

    # 章节识别模式
    SECTION_PATTERNS = {
        "1": r"^1[、.]\s*发明创造名称",
        "2": r"^2[、.]\s*所属技术领域",
        "3": r"^3[、.]\s*相关的背景技术",
        "4": r"^4[、.]\s*发明内容",
        "4.1": r"^（1）[、.]\s*解决的技术问题",
        "4.2": r"^（2）[、.]\s*技术方案",
        "4.3": r"^（3）[、.]\s*有益效果",
        "5": r"^5[、.]\s*具体实施方式",
        "6": r"^6[、.]\s*关键点和欲保护点",
        "7": r"^7[、.]\s*其他有助于理解本技术的资料",
    }

    def __init__(self, template_path: str, output_path: str):
        """
        初始化生成器

        Args:
            template_path: DOCX 模板文件路径
            output_path: 输出 DOCX 文件路径
        """
        self.template_path = Path(template_path)
        self.output_path = Path(output_path)

        if not self.template_path.exists():
            raise TemplateNotFoundError(str(self.template_path))

        # 检查字体
        self.font_checker = FontChecker()
        if not self.font_checker.is_font_available(self.FONT_NAME):
            raise FontNotFoundError(
                self.FONT_NAME, FontInstaller.get_installation_guide()
            )

        self.doc = None
        self.formatter = ParagraphFormatter()

    def load_template(self):
        """加载 DOCX 模板"""
        self.doc = Document(str(self.template_path))

    def load_parsed_data(self, json_path: str) -> Dict[str, Any]:
        """
        加载解析后的 JSON 数据

        Args:
            json_path: JSON 文件路径

        Returns:
            解析后的数据字典
        """
        with open(json_path, "r", encoding="utf-8") as f:
            return json.load(f)

    def fill_content(self, parsed_data: Dict[str, Any]):
        """
        填充内容到模板

        Args:
            parsed_data: 解析后的章节数据
        """
        sections = parsed_data.get("sections", [])

        # 构建章节编号到内容的映射
        section_map = {}
        for section in sections:
            number = section["number"]
            content = section.get("content", "")

            # 如果有子章节，分别处理
            if "subsections" in section:
                for subsection in section["subsections"]:
                    sub_number = subsection["number"]
                    section_map[sub_number] = subsection.get("content", "")
            else:
                section_map[number] = content

        # 遍历模板段落，填充内容
        current_section = None

        for para in self.doc.paragraphs:
            text = para.text.strip()

            # 识别当前章节
            identified_section = self._identify_section(text)
            if identified_section:
                current_section = identified_section
                continue

            # 如果是备注段落（包含【】），则替换为实际内容
            if current_section and "【" in text and "】" in text:
                content = section_map.get(current_section, "")
                if content:
                    self._fill_paragraph(para, content, current_section)

    def _identify_section(self, text: str) -> str:
        """
        识别段落所属章节

        Args:
            text: 段落文本

        Returns:
            章节编号，如果未识别则返回 None
        """
        for section_num, pattern in self.SECTION_PATTERNS.items():
            if re.match(pattern, text):
                return section_num
        return None

    def _fill_paragraph(self, para, content: str, section_num: str):
        """
        填充段落内容并应用格式

        Args:
            para: 目标段落对象
            content: 要插入的内容
            section_num: 章节编号
        """
        # 清空原段落
        para.clear()

        # 判断是标题还是正文
        is_title = section_num in ["1"]  # 第1章节是标题

        # 分段内容（按段落标记）
        paragraphs = content.split("\n\n")

        # 插入第一段
        run = para.add_run(paragraphs[0] if paragraphs else "")
        if is_title:
            self._apply_title_style(run)
        else:
            self._apply_body_style(run)

        # 插入后续段落
        for p_text in paragraphs[1:]:
            if p_text.strip():
                new_para = para._element.addnext(
                    para._element.__class__()
                )
                new_para_obj = Document(new_para.getparent()).paragraphs[-1]
                run = new_para_obj.add_run(p_text)
                self._apply_body_style(run)
                self.formatter.format_body_paragraph(new_para_obj)

    def _apply_title_style(self, run):
        """应用标题字体样式"""
        run.font.name = self.TITLE_FONT_NAME
        run.font.size = Pt(self.TITLE_FONT_SIZE)
        run.font.bold = True

        # 设置中文字体（必需）
        run._element.rPr.rFonts.set(
            "{http://schemas.openxmlformats.org/drawingml/2006/main}eastAsia",
            self.TITLE_FONT_NAME,
        )

    def _apply_body_style(self, run):
        """应用正文字体样式"""
        run.font.name = self.BODY_FONT_NAME
        run.font.size = Pt(self.BODY_FONT_SIZE)
        run.font.bold = False

        # 设置中文字体（必需）
        run._element.rPr.rFonts.set(
            "{http://schemas.openxmlformats.org/drawingml/2006/main}eastAsia",
            self.BODY_FONT_NAME,
        )

    def format_document(self):
        """格式化整个文档"""
        # 设置页面
        self.formatter.format_page_setup(self.doc)

        # 遍历所有段落，设置格式
        for para in self.doc.paragraphs:
            text = para.text.strip()

            # 跳过空段落
            if not text:
                continue

            # 判断是否是章节标题
            if self._identify_section(text):
                # 章节标题保持原有格式，只设置字体
                if para.runs:
                    for run in para.runs:
                        if run.text.strip():
                            self._apply_title_style(run)
                self.formatter.format_heading_paragraph(para)
            # 正文字段（已填充内容）
            elif para.runs:
                for run in para.runs:
                    if run.text.strip():
                        # 确保字体正确
                        self._apply_body_style(run)
                self.formatter.format_body_paragraph(para)

    def save(self):
        """保存生成的 DOCX 文件"""
        self.output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(self.output_path))

    def generate(self, json_path: str) -> Dict[str, Any]:
        """
        完整生成流程

        Args:
            json_path: 解析后的 JSON 文件路径

        Returns:
            生成统计信息
        """
        # 加载模板
        self.load_template()

        # 加载数据
        parsed_data = self.load_parsed_data(json_path)

        # 填充内容
        self.fill_content(parsed_data)

        # 格式化文档
        self.format_document()

        # 保存
        self.save()

        # 统计信息
        stats = {
            "success": True,
            "docx_path": str(self.output_path),
            "generation_timestamp": datetime.now().isoformat(),
            "stats": {
                "total_paragraphs": len(self.doc.paragraphs),
                "sections_filled": len(parsed_data.get("sections", [])),
                "font_applied": {
                    "title": self.TITLE_FONT_NAME,
                    "body": self.BODY_FONT_NAME,
                },
            },
        }

        return stats


def main():
    """命令行入口"""
    import sys

    if len(sys.argv) < 4:
        print(
            "Usage: python docx_generator.py <json_path> <template_path> <output_path>"
        )
        sys.exit(1)

    json_path = sys.argv[1]
    template_path = sys.argv[2]
    output_path = sys.argv[3] if len(sys.argv) > 3 else "output.docx"

    try:
        generator = DocxGenerator(template_path, output_path)
        stats = generator.generate(json_path)

        print(f"✅ 生成成功！")
        print(f"📄 输出文件: {stats['docx_path']}")
        print(f"📊 总段落数: {stats['stats']['total_paragraphs']}")
        print(f"📝 填充章节: {stats['stats']['sections_filled']}")
        print(f"🎨 字体应用:")
        print(f"   标题: {stats['stats']['font_applied']['title']}")
        print(f"   正文: {stats['stats']['font_applied']['body']}")

    except (FontNotFoundError, TemplateNotFoundError) as e:
        print(f"❌ {e}")
        sys.exit(1)
    except Exception as e:
        print(f"❌ 生成失败: {e}")
        import traceback

        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
