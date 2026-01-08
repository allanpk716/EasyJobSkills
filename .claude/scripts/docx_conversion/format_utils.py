"""
格式工具函数 - 设置段落和文档格式
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


class ParagraphFormatter:
    """段落格式化器"""

    # 页面设置（专利标准）
    PAGE_MARGINS = {
        "top": 1.0,  # 1英寸 = 2.54cm
        "bottom": 1.0,
        "left": 1.25,  # 1.25英寸 = 3.17cm
        "right": 1.25,
    }

    # 段落格式
    LINE_SPACING = 1.5
    FIRST_LINE_INDENT = Inches(0.25)  # 约2个字符

    def format_page_setup(self, doc: Document):
        """
        设置页面格式

        Args:
            doc: Document 对象
        """
        section = doc.sections[0]

        # 设置页边距
        section.top_margin = Inches(self.PAGE_MARGINS["top"])
        section.bottom_margin = Inches(self.PAGE_MARGINS["bottom"])
        section.left_margin = Inches(self.PAGE_MARGINS["left"])
        section.right_margin = Inches(self.PAGE_MARGINS["right"])

        # 设置纸张大小为 A4
        section.page_width = Inches(8.27)  # A4 宽度
        section.page_height = Inches(11.69)  # A4 高度

    def format_heading_paragraph(self, para):
        """
        格式化标题段落

        Args:
            para: Paragraph 对象
        """
        # 设置对齐
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # 设置行距
        para.paragraph_format.line_spacing = 1.0

        # 清除首行缩进
        para.paragraph_format.first_line_indent = None

    def format_body_paragraph(self, para):
        """
        格式化正文段落

        Args:
            para: Paragraph 对象
        """
        # 设置对齐
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # 设置行距
        para.paragraph_format.line_spacing = self.LINE_SPACING

        # 设置首行缩进
        para.paragraph_format.first_line_indent = self.FIRST_LINE_INDENT

        # 设置段落间距
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)


if __name__ == "__main__":
    # 测试代码
    from docx import Document

    doc = Document()
    formatter = ParagraphFormatter()

    # 设置页面格式
    formatter.format_page_setup(doc)
    print("页面格式设置完成")

    # 添加标题段落
    heading = doc.add_paragraph("这是一个标题")
    formatter.format_heading_paragraph(heading)
    print("标题段落格式化完成")

    # 添加正文段落
    body = doc.add_paragraph("这是正文内容，用于测试段落格式化。")
    formatter.format_body_paragraph(body)
    print("正文段落格式化完成")

    # 保存测试文档
    test_path = "test_format.docx"
    doc.save(test_path)
    print(f"测试文档已保存: {test_path}")
